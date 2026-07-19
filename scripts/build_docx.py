#!/usr/bin/env python3
"""Build career-destination Word report from JSON content.

Usage:
  python3 build_docx.py --major Marketing --major-zh 市场营销 \\
    --regions "北美,中国大陆" --out report.docx --json content.json
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    print("Missing python-docx. Install: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


def set_run_font(run, size=11, bold=False, color=None, name_cn="宋体", name_en="Calibri"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_cn)
    if color:
        run.font.color.rgb = color


def add_heading_cn(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=20, bold=True, name_cn="黑体", color=RGBColor(0x1A, 0x1A, 0x2E))
        p.paragraph_format.space_after = Pt(6)
    elif level == 1:
        run = p.add_run(text)
        set_run_font(run, size=15, bold=True, name_cn="黑体", color=RGBColor(0x1A, 0x3A, 0x5C))
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    else:
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, name_cn="黑体", color=RGBColor(0x2C, 0x5F, 0x8A))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11, name_cn="宋体")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=11, name_cn="宋体")
    p.paragraph_format.space_after = Pt(3)
    return p


def shade_header_row(row, color="1A3A5C"):
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, name_cn="黑体", color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_header_row(hdr)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=9.5, name_cn="宋体")
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()
    return table


def build_report(data: dict, out: Path) -> None:
    major = data.get("major") or "Major"
    major_zh = data.get("major_zh") or major
    regions = data.get("regions") or []
    if isinstance(regions, str):
        regions = [r.strip() for r in regions.split(",") if r.strip()]

    date = data.get("date") or datetime.now().strftime("%Y年%m月%d日")
    if "-" in str(date) and "年" not in str(date):
        try:
            dt = datetime.strptime(str(date)[:10], "%Y-%m-%d")
            date = dt.strftime("%Y年%m月%d日")
        except ValueError:
            pass

    title_major = major_zh if major_zh else major
    regions_label = " / ".join(regions) if regions else "未指定地区"

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    add_heading_cn(doc, f"{title_major}专业就业去向与顶尖用人企业汇总", 0)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"整理日期：{date}  |  覆盖范围：{regions_label}")
    set_run_font(run, size=10, name_cn="楷体", color=RGBColor(0x66, 0x66, 0x66))

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "说明：综合官方劳工统计、行业招聘报告与高校/企业公开信息整理，供升学与求职参考。"
        "雇主列表为高频目标，非官方完整排名。"
    )
    set_run_font(run, size=9, name_cn="宋体", color=RGBColor(0x88, 0x88, 0x88))

    add_heading_cn(doc, "一、专业概况与就业前景", 1)
    outlook = data.get("outlook") or {}
    if outlook.get("summary"):
        add_body(doc, outlook["summary"])
    for b in outlook.get("bullets") or []:
        add_bullet(doc, b)
    for s in outlook.get("stats") or []:
        label, value, note_s = s.get("label", ""), s.get("value", ""), s.get("note", "")
        add_bullet(doc, f"{label}：{value}" + (f"（{note_s}）" if note_s else ""))

    add_heading_cn(doc, "二、主要就业方向", 1)
    paths = data.get("paths") or []
    if paths:
        add_table(
            doc,
            ["就业方向", "典型岗位", "能力侧重", "常见行业"],
            [
                [
                    p.get("name", ""),
                    p.get("roles", ""),
                    p.get("skills", ""),
                    p.get("industries", ""),
                ]
                for p in paths
            ],
        )
    else:
        add_body(doc, "（本次未提供结构化就业方向表。）")

    add_heading_cn(doc, "三、典型职业路径", 1)
    ladders = data.get("ladders") or []
    if ladders:
        for ladder in ladders:
            steps = ladder.get("steps") or []
            add_bullet(doc, f"{ladder.get('title', '')}：{' → '.join(steps)}")
    else:
        add_body(doc, "（见补充的晋升路径说明。）")

    add_heading_cn(doc, "四、顶尖用人企业（分地区）", 1)
    employers = data.get("employers_by_region") or {}
    region_order = regions if regions else list(employers.keys())
    for idx, region in enumerate(region_order, start=1):
        add_heading_cn(doc, f"4.{idx} {region}", 2)
        groups = employers.get(region) or []
        if not groups:
            add_body(doc, f"（暂无 {region} 雇主分组数据。）")
            continue
        for g in groups:
            add_bullet(doc, f"【{g.get('group', '其他')}】")
            for item in g.get("items") or []:
                name = item.get("name", "")
                roles = item.get("roles", "")
                badge = item.get("badge", "")
                line = f"{name} — {roles}" if roles else name
                if badge:
                    line += f" 〔{badge}〕"
                add_bullet(doc, line)

    compare = data.get("compare") or []
    if len(region_order) >= 2 and compare:
        add_heading_cn(doc, "五、地区对比要点", 1)
        headers = ["维度"] + region_order
        rows = []
        for row in compare:
            cells_map = row.get("cells") or {}
            rows.append(
                [row.get("dimension", "")] + [cells_map.get(r, "") for r in region_order]
            )
        add_table(doc, headers, rows)
        prep_heading = "六、求职准备建议"
        src_heading = "七、主要参考来源"
    else:
        prep_heading = "五、求职准备建议"
        src_heading = "六、主要参考来源"

    add_heading_cn(doc, prep_heading, 1)
    skills = data.get("skills") or {}
    add_heading_cn(doc, "硬技能", 2)
    for x in skills.get("hard") or []:
        add_bullet(doc, x)
    add_heading_cn(doc, "软技能", 2)
    for x in skills.get("soft") or []:
        add_bullet(doc, x)
    add_heading_cn(doc, "趋势与差异化", 2)
    for x in skills.get("edge") or []:
        add_bullet(doc, x)

    add_heading_cn(doc, src_heading, 1)
    for s in data.get("sources") or []:
        add_bullet(doc, s)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— 文档结束 —")
    set_run_font(run, size=10, name_cn="宋体", color=RGBColor(0x99, 0x99, 0x99))

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Saved: {out}")


def main():
    ap = argparse.ArgumentParser(description="Build career destination DOCX")
    ap.add_argument("--major", required=True)
    ap.add_argument("--major-zh", default="")
    ap.add_argument("--regions", default="", help="Comma-separated regions")
    ap.add_argument("--out", required=True)
    ap.add_argument("--json", required=True, help="Path to content.json")
    args = ap.parse_args()

    path = Path(args.json)
    data = json.loads(path.read_text(encoding="utf-8"))
    data["major"] = args.major or data.get("major")
    data["major_zh"] = args.major_zh or data.get("major_zh") or data["major"]
    if args.regions:
        data["regions"] = [r.strip() for r in args.regions.split(",") if r.strip()]

    build_report(data, Path(args.out))


if __name__ == "__main__":
    main()
